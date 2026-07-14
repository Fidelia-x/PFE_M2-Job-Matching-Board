# test_cv_parser.py
# Teste l'extraction de texte des CV (PDF/DOCX) et, en option, le matching
# complet, sans passer par l'interface Streamlit.
#
# Usage :
#   python back_service/test_cv_parser.py chemin/vers/cv.pdf
#   python back_service/test_cv_parser.py chemin/vers/cv.docx --matching
#
# Sans argument, lance un test avec un PDF et un DOCX générés en mémoire
# (aucun fichier réel requis) pour vérifier que l'extraction fonctionne.

import os
import sys
import argparse

# Quand ce fichier est lancé directement (python back_service/test_cv_parser.py),
# Python met back_service/ sur sys.path, pas frontend/ — du coup le package
# back_service lui-même n'est pas trouvable. On ajoute frontend/ explicitement.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from back_service.cv_parser import extract_cv_text

class FakeUploadedFile:
    """Reproduit l'interface attendue par extract_cv_text (.getvalue(), .name)
    à partir d'un fichier sur disque, sans dépendre de Streamlit."""
    def __init__(self, path):
        self.name = path
        with open(path, "rb") as f:
            self._data = f.read()

    def getvalue(self):
        return self._data


def test_extraction_from_path(cv_path):
    print(f"--- Extraction : {cv_path} ---")
    text = extract_cv_text(FakeUploadedFile(cv_path))
    print(f"Longueur du texte extrait : {len(text)} caractères")
    print(text[:500] + ("..." if len(text) > 500 else ""))
    return text


def test_matching(cv_text, top_n=5):
    from scripts.matching_cv import find_best_matches
    print(f"\n--- Matching (top {top_n}) ---")
    matches = find_best_matches(cv_text, top_n=top_n)
    if not matches:
        print("Aucun match trouvé. Vérifiez que 'offres_emploi' contient des données avec embeddings.")
        return
    for m in matches:
        print(f"{m['titre']} — {m['company']} ({round(m['score'] * 100)}%)")


def test_extraction_sample_files():
    """Génère un PDF et un DOCX de test en mémoire pour vérifier que
    l'extraction fonctionne, sans avoir besoin d'un vrai CV sous la main."""
    import io

    print("--- Test avec un DOCX généré en mémoire ---")
    from docx import Document
    doc = Document()
    doc.add_paragraph("Jean Dupont")
    doc.add_paragraph("Développeur Python expert en Airflow et SQL, cherche poste de Data Engineer.")
    buf = io.BytesIO()
    doc.save(buf)

    class FakeUploadedBytes:
        def __init__(self, data, name):
            self._data = data
            self.name = name
        def getvalue(self):
            return self._data

    text = extract_cv_text(FakeUploadedBytes(buf.getvalue(), "fake.docx"))
    assert "Python" in text and "Data Engineer" in text, "Extraction DOCX incorrecte"
    print("OK :", text.replace("\n", " | "))

    print("\n--- Test avec un PDF généré en mémoire ---")
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Jean Dupont")
    c.drawString(100, 730, "Developpeur Python expert en Airflow et SQL cherche poste Data Engineer.")
    c.save()

    text = extract_cv_text(FakeUploadedBytes(buf.getvalue(), "fake.pdf"))
    assert "Python" in text and "Data Engineer" in text, "Extraction PDF incorrecte"
    print("OK :", text.replace("\n", " | "))

    print("\n--- Test format non supporté ---")
    try:
        extract_cv_text(FakeUploadedBytes(b"peu importe", "fake.txt"))
        print("ERREUR : aurait dû lever une exception")
    except ValueError as e:
        print("OK, rejeté comme attendu :", e)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Teste l'extraction de CV (et en option le matching), sans passer par Streamlit.")
    parser.add_argument("cv_path", nargs="?", help="Chemin vers un CV réel (PDF/DOCX). Sans argument : test avec des fichiers générés en mémoire.")
    parser.add_argument("--matching", action="store_true", help="Enchaîne aussi le matching réel (nécessite la DB + le modèle)")
    parser.add_argument("--top", type=int, default=5, help="Nombre d'offres à afficher")
    args = parser.parse_args()

    if args.cv_path:
        try:
            cv_text = test_extraction_from_path(args.cv_path)
        except ValueError as e:
            print(f"Erreur d'extraction : {e}")
            sys.exit(1)

        if args.matching:
            try:
                test_matching(cv_text, top_n=args.top)
            except Exception as e:
                print(f"Erreur lors du matching : {e}")
                sys.exit(1)
    else:
        test_extraction_sample_files()
